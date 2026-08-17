# netsage_agent.py
# Assistant Name: NetSage AI (Ollama Powered)
#
# Two ways to run this file:
#   python netsage_agent.py            -> CLI mode (same as before)
#   python netsage_agent.py --web      -> Web server on http://0.0.0.0:5001
#
# The web server is its OWN Flask app, separate from your main app.py
# (which should keep running on :5000). This file talks to app.py over
# HTTP via ORCHESTRATOR_URL, exactly like the CLI version always did.

import os
import sys
import glob
import re
import time
import threading
import json
import requests
import ollama
from pathlib import Path
from flask import Flask, render_template, request, jsonify, Response, send_from_directory

# ---------------------------------------------------------
# Config
# ---------------------------------------------------------
ORCHESTRATOR_URL = os.environ.get("ORCHESTRATOR_URL", "http://localhost:5000/api/action")
OLLAMA_MODEL = os.environ.get("OLLAMA_MODEL", "llama3.1")

PROMETHEUS_URL = os.environ.get("PROMETHEUS_URL", "http://localhost:9090").rstrip("/")
GRAFANA_URL = os.environ.get("GRAFANA_URL", "http://localhost:3000").rstrip("/")
# Create a Grafana Service Account token (Administration > Service accounts)
# and export it as GRAFANA_API_KEY before starting this script, or Grafana
# dashboard/KPI lookups will fail with a 401 (Prometheus itself needs no auth).
GRAFANA_API_KEY = os.environ.get("GRAFANA_API_KEY", "")

REPORTS_DIR = Path("reports")
REPORTS_DIR.mkdir(parents=True, exist_ok=True)

MAX_TOOL_ITERATIONS = 5  # lets the agent chain tool calls, e.g. search a
                          # Grafana dashboard, then query Prometheus with
                          # what it found, before giving a final answer.

# --------------------------------------------------------------------
# NEW: Helper to clean text for XML compatibility (used in Word reports)
# --------------------------------------------------------------------
def clean_text(text: str) -> str:
    """Remove null bytes and control characters (except newline/tab) 
    to make text safe for python-docx (XML)."""
    if not isinstance(text, str):
        text = str(text)
    # Remove all control chars except \n (0x0A) and \t (0x09)
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)

# --------------------------------------------------------------------
# UPDATED SYSTEM PROMPT – now includes explicit rule for tool errors
# --------------------------------------------------------------------
SYSTEM_PROMPT = (
    "You are NetSage AI, a warm, senior, and conversational 5G network core engineer. "
    "CRITICAL ARCHITECTURE RULE: This is a 100% pure 5G Standalone (SA) Core network using "
    "Open5GS components (AMF, SMF, UPF, UDM, UDR, AUSF, NRF) and UERANSIM. "
    "Never mention 4G/EPC components like HSS or MME, as they do not exist in this environment. "
    "CRITICAL EXECUTION RULE: NEVER hallucinate, simulate, or make up JSON output, tool responses, "
    "container statuses, or KPI numbers. If you need container health, call `check_container_health`. "
    "If you need a live metric, call `query_prometheus`. Never invent a number. "
    "ABSOLUTE RULE — NO FAKE TOOL CALLS IN TEXT: after your real tools have run, you are DONE calling "
    "tools for this turn. Do NOT write code blocks, backticks, or prose that look like a tool invocation "
    "(e.g. 'check_container_health' followed by a made-up output, or 'query_prometheus ...' followed by "
    "invented numbers). If you have not actually received a tool result for something in this turn, do "
    "not mention a value for it at all — simply don't discuss it, or say you'd need to check and offer to "
    "do so next. Only report data that came back from a tool that was genuinely called this turn. "
    "ABSOLUTE RULE — NEVER PARAPHRASE NAMES OR IDENTIFIERS: metric names, container names, file names, "
    "IMSIs, IPs, and any other exact identifier MUST be copied character-for-character from the tool "
    "output. Never invent a plausible-looking name, never 'clean up' or shorten an identifier, never guess "
    "at a naming convention. If `list_prometheus_metrics` returns real names, quote them exactly as given "
    "— do not summarize them into names that sound right but weren't actually in the output. If you're "
    "not sure a name is exactly right, say so rather than smoothing it over. "
    "CRITICAL KPI RULE: If you don't know the exact PromQL metric name for something the user asks about, "
    "first call `list_prometheus_metrics` (with a filter like 'amf', 'ue', 'ran') or "
    "`search_grafana_dashboards` + `get_grafana_dashboard_kpis` to find the real expression used on the "
    "dashboards, THEN call `query_prometheus` with that exact expression. Report the tool's exact output; "
    "don't reformat numbers speculatively. If `query_prometheus` comes back with no data, do not just "
    "report the failure — automatically call `list_prometheus_metrics` with a relevant filter to find the "
    "correct metric name, then retry `query_prometheus` with it, before giving your final answer. Only "
    "explain the limitation to the user if that retry also fails. "
    "CRITICAL IDENTITY/CONFIG RULE: Values like IMSI, K/OPc keys, APN, or other subscriber identity fields "
    "are configured values, NOT something that appears in runtime container logs. If asked for a UE's IMSI "
    "or similar config detail, call `read_yaml_file` on that UE's config (e.g. 'nr-ue1.yaml') instead of "
    "searching logs for it. "
    "CRITICAL TOPOLOGY RULE: If the user asks to see, visualize, or get a map/diagram/overview of the "
    "network topology or layout, call `get_network_topology`. The web UI renders that data as a live "
    "diagram automatically — your reply text must be ONE short sentence introducing it, nothing more. Do "
    "not describe the topology yourself in ASCII art, bullet lists, or long prose, and do not add "
    "unrelated suggestions (e.g. recommending third-party visualization tools) that weren't asked for. "
    "CRITICAL DEPTH RULE: Outside of the topology diagram case above, don't just state a bare result — "
    "briefly explain what it means. E.g. not just 'AMF registered UE count: 12' but a sentence on what "
    "that indicates. Every bit of explanation must be grounded in the real data you were actually given "
    "this turn — elaborate on real numbers and real log lines, never invent extra specifics, causes, or "
    "context you weren't actually given just to sound more thorough. "
    "CRITICAL REPORT RULE: If the user asks for a network report, a summary document, or something to "
    "download/save/share about the current network state, call `generate_network_report`. The web UI "
    "renders a clickable download button for the file automatically — do not write out the URL yourself "
    "in your reply (you're prone to mistyping or dropping it). Just give a one or two sentence summary of "
    "the key numbers from the tool result. "
    "CRITICAL DEPENDENCY RULE: The Core network MUST always be started and fully initialized FIRST "
    "before any gNodeBs (gnb) or UEs can be deployed. Never deploy gNBs or UEs while the core is down. "
    "Always execute deployment steps sequentially: 1) Start Core -> 2) Start gNBs -> 3) Add UEs. "
    "CRITICAL PARAMETER RULE: If the user gives an unclear or invalid number (typo, garbled text, etc.) "
    "for how many gNBs or UEs to deploy, do NOT guess or silently substitute a default. Ask them to confirm "
    "the exact number before calling any deployment tool. "
    "CRITICAL COMMUNICATION & LOG SIMPLIFICATION RULE: "
    "1. Always talk like a friendly, clear, and encouraging human engineer. Never sound robotic. "
    "2. Whenever you look up logs, never dump raw text straight out. Summarize key milestones into neat bullet points. "
    "3. Never invent errors or reference non-existent components (like AMF[2] or 4G entities). If logs or "
    "metrics are normal, explicitly state they are normal. "
    # NEW RULE – explicitly forbids fabrication when a tool fails
    "CRITICAL ERROR HANDLING RULE: If a tool call returns an error or returns no data (empty result), "
    "you MUST NOT invent or guess any numbers, statuses, or conclusions. You must clearly state that the "
    "tool returned an error and offer to retry or to check individual components. Never generate a summary "
    "that includes values not actually returned by a successful tool call. For example, if generate_network_report "
    "fails, say 'The report generation failed due to [error message]' and do not add any fabricated KPIs or "
    "summary statistics. If you don't have real data, say so and ask what the user would like to do next."
)

# ---------------------------------------------------------
# Spinner Utility (CLI mode only)
# ---------------------------------------------------------
class Spinner:
    def __init__(self, message="Thinking..."):
        self.message = message
        self.stop_running = False
        self.thread = None

    def spinner_task(self):
        chars = "⠋⠙⠹⠸⠼⠴⠦⠧⠇⠏"
        i = 0
        while not self.stop_running:
            sys.stdout.write(f"\r[⏳ NetSage AI] {chars[i % len(chars)]} {self.message}")
            sys.stdout.flush()
            time.sleep(0.1)
            i += 1
        sys.stdout.write("\r" + " " * (len(self.message) + 20) + "\r")
        sys.stdout.flush()

    def __enter__(self):
        self.stop_running = False
        self.thread = threading.Thread(target=self.spinner_task)
        self.thread.start()
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.stop_running = True
        if self.thread:
            self.thread.join()

# ---------------------------------------------------------
# Orchestrator tools (talk to app.py on :5000 over HTTP)
# ---------------------------------------------------------
def check_container_health(component: str = "core", **kwargs) -> str:
    """Check the running status and health of the 5G network: core state, and
    every gNB/UE/AMF container currently up, using the orchestrator's live
    status endpoint (the same data source the dashboard polls).
    Args:
        component: Optional filter, currently informational only (e.g. 'core', 'gnb', 'ue', 'amf')
    """
    try:
        resp = requests.get("http://localhost:5000/api/status", timeout=8)
        if not resp.ok:
            return f"Status check failed: {resp.status_code} {resp.text[:200]}"
        data = resp.json()
    except Exception as e:
        return f"Failed to reach orchestrator status endpoint: {e}"

    core_state = data.get("core", "UNKNOWN")
    gnbs = data.get("gnbs", {})
    ues = data.get("ues", {})
    amfs = data.get("amfs", {})

    gnb_running = [k for k, v in gnbs.items() if v.get("status") == "running"]
    ue_running = [k for k, v in ues.items() if v.get("status") == "running"]
    amf_running = [k for k, v in amfs.items() if v.get("status") == "running"]
    amf_errored = [k for k, v in amfs.items() if v.get("status") != "running"]

    def _sorted(ids):
        return sorted(ids, key=lambda x: int(x)) if ids else []

    lines = [
        f"Core: {core_state}",
        f"gNBs running: {len(gnb_running)} (ids: {', '.join(_sorted(gnb_running)) or 'none'})",
        f"UEs running: {len(ue_running)} (ids: {', '.join(_sorted(ue_running)) or 'none'})",
        f"AMF instances up: {len(amf_running)} (ids: {', '.join(_sorted(amf_running)) or 'none'})",
    ]
    if amf_errored:
        lines.append(
            f"AMF instances with errors: {', '.join(_sorted(amf_errored))} "
            f"(failed to start correctly, check container logs for these)"
        )

    return "\n".join(lines)

def start_core(**kwargs) -> str:
    """Start the 5G core network infrastructure only if not already running."""
    health_status = check_container_health()
    if "core: online" in health_status.lower():
        return "Core is already running. Skipping start step."
    return execute_action("start_core")

def launch_full_network(num_gnbs: int = 2, ues_per_gnb: int = 2, **kwargs) -> str:
    """Launch the entire 5G network sequentially: Start Core -> Start gNBs -> Add UEs."""
    results = []
    results.append(start_core())
    time.sleep(5)
    for gnb_id in range(1, int(num_gnbs) + 1):
        results.append(start_gnb(gnb_id=gnb_id, tac=gnb_id))
        time.sleep(3)
        for u in range(1, int(ues_per_gnb) + 1):
            ue_id = (gnb_id - 1) * int(ues_per_gnb) + u
            results.append(add_ue(ue_id=ue_id, target_gnb=gnb_id))
            time.sleep(2)
    return " | ".join(results)

def stop_core() -> str:
    """Stop and tear down the 5G core network."""
    return execute_action("stop_core")

def start_gnb(gnb_id: int, tac: int = 1) -> str:
    """Deploy a new gNodeB (RAN simulator).
    Args:
        gnb_id: The ID of the gNodeB (e.g., 1, 2)
        tac: Tracking Area Code for this gNodeB (default is 1)
    """
    return execute_action("start_gnb", {"gnb_id": int(gnb_id), "tac": int(tac)})

def add_ue(ue_id: int, target_gnb: int) -> str:
    """Deploy a User Equipment (UE) connected to a target gNB.
    Args:
        ue_id: The ID of the UE (e.g., 1, 2)
        target_gnb: The gNB ID this UE should connect to
    """
    return execute_action("add_ue", {"ue_id": int(ue_id), "target_gnb": int(target_gnb)})

def get_ue_logs(ue_id: int = 1) -> str:
    """Get the live Docker container logs for a specific User Equipment (UE).
    Args:
        ue_id: The ID of the UE (e.g., 1, 2)
    """
    return _fetch_container_logs(f"nr_ue{int(ue_id)}")

def get_gnb_logs(gnb_id: int = 1) -> str:
    """Get the live Docker container logs for a specific gNodeB.
    Args:
        gnb_id: The ID of the gNodeB (e.g., 1, 2)
    """
    return _fetch_container_logs(f"nr_gnb{int(gnb_id)}")

def get_core_logs(component: str = "amf") -> str:
    """Get the live Docker container logs for a core network component (e.g., amf, smf, upf, ausf, udm, udr, nrf, pcf).
    Args:
        component: The core network component name (default is amf)
    """
    return _fetch_container_logs(str(component))

def _fetch_container_logs(container_name: str) -> str:
    try:
        response = requests.get(
            "http://localhost:5000/api/container-logs",
            params={"container": container_name},
            timeout=5
        )
        if response.ok:
            data = response.json()
            logs = data.get("logs", "")
            return logs if logs.strip() else "CONTAINER_LOGS_EMPTY"
        return f"Error: Received status {response.status_code}"
    except Exception as e:
        return f"Failed to connect to orchestrator API: {e}"

def stop_all(**kwargs) -> str:
    """Trigger full cleanup, removing all containers and generated files."""
    return execute_action("stop_all")

def get_logs_or_files(target_name: str = "list") -> str:
    """Inspect local YAML configuration files, report files, logs, or list all available YAML configs in the workspace.
    Args:
        target_name: The filename to inspect (e.g., 'nr-gnb.yaml', 'smf.yaml'), or type 'list' to see all YAML files.
    """
    try:
        if target_name.lower() in ["list", "all", "yamls", "files"]:
            yaml_files = glob.glob("*.yaml")
            return f"Available YAML configuration files in workspace: {yaml_files}"
        if os.path.exists(target_name) and os.path.isfile(target_name):
            with open(target_name, 'r') as f:
                return f.read()[:4000]
        payload = {"action": "get_logs", "target": target_name}
        response = requests.post(ORCHESTRATOR_URL, json=payload, timeout=10)
        if response.ok:
            return str(response.json())
        return f"Could not find local file or log for '{target_name}'. Use 'list' to see available YAML files."
    except Exception as e:
        return f"Error reading target {target_name}: {e}"

def read_yaml_file(target_name: str) -> str:
    """Read the contents of a YAML configuration file.
    Args:
        target_name: The filename to read (e.g., 'smf.yaml', 'amf.yaml', 'nr-gnb.yaml')
    """
    try:
        if os.path.exists(target_name) and os.path.isfile(target_name):
            with open(target_name, 'r') as f:
                return f.read()[:4000]
        return f"File '{target_name}' not found."
    except Exception as e:
        return f"Error reading file: {e}"

def get_network_topology(**kwargs) -> str:
    """Get a structured summary of the current network topology: core state,
    core network functions, AMF instance health, gNodeB count/TACs, and UE
    count. Call this whenever the user asks to see, visualize, or get a map,
    diagram, or overview of the network topology or layout. The web UI
    renders this data as a live diagram automatically — just briefly
    introduce it in your reply text, don't describe the topology yourself
    in ASCII art or a long list.
    """
    try:
        resp = requests.get("http://localhost:5000/api/status", timeout=8)
        if not resp.ok:
            return f"Status check failed: {resp.status_code} {resp.text[:200]}"
        data = resp.json()
    except Exception as e:
        return f"Failed to reach orchestrator status endpoint: {e}"

    core_state = data.get("core", "UNKNOWN")
    gnbs = data.get("gnbs", {})
    ues = data.get("ues", {})
    amfs = data.get("amfs", {})

    gnb_running = {k: v for k, v in gnbs.items() if v.get("status") == "running"}
    ue_running = {k: v for k, v in ues.items() if v.get("status") == "running"}
    amf_running = [k for k, v in amfs.items() if v.get("status") == "running"]
    amf_error = [k for k, v in amfs.items() if v.get("status") != "running"]
    tacs = sorted([v.get("tac") for v in gnb_running.values() if v.get("tac") is not None])

    topology = {
        "core_state": core_state,
        "functions": ["AMF", "SMF", "UPF", "UDM", "UDR", "AUSF", "NRF"],
        "amf_running": len(amf_running),
        "amf_error": len(amf_error),
        "gnb_count": len(gnb_running),
        "gnb_tacs": tacs,
        "ue_count": len(ue_running),
    }
    return json.dumps(topology)

LOG_KEYWORDS_UE = ["registration", "pdu session", "psi", "deregistration", "handover", "failed", "error"]
LOG_KEYWORDS_GNB = ["ng setup", "sctp", "ue context", "registered", "connected", "failed", "error"]
MAX_DEVICES_FOR_LOG_SUMMARY = 20

def _summarize_log(raw_log, keywords, max_lines=4):
    """Pull out milestone lines by keyword match — deterministic, no LLM
    involved, so the summary can't drift from what the log actually says."""
    if not raw_log or raw_log == "CONTAINER_LOGS_EMPTY":
        return []
    matched = [ln.strip() for ln in raw_log.splitlines() if any(kw in ln.lower() for kw in keywords)]
    return matched[-max_lines:] if matched else []

def generate_network_report(**kwargs) -> str:
    """Generate a downloadable Word (.docx) report of the current network:
    core state and functions, AMF instance health, gNodeBs, connected UEs,
    Prometheus scrape target health, and recent errors pulled from the real
    deployment console log. Call this whenever the user asks for a network
    report, a summary document, or anything to download/save/share about
    the current network state.
    """
    try:
        from docx import Document
    except ImportError:
        return (
            "The python-docx library isn't installed on the server. Run this in the "
            "venv and try again: pip install python-docx"
        )

    try:
        resp = requests.get("http://localhost:5000/api/status", timeout=8)
        if not resp.ok:
            return f"Status check failed: {resp.status_code} {resp.text[:200]}"
        data = resp.json()
    except Exception as e:
        return f"Failed to reach orchestrator status endpoint: {e}"

    core_state = data.get("core", "UNKNOWN")
    gnbs = data.get("gnbs", {})
    ues = data.get("ues", {})
    amfs = data.get("amfs", {})
    console_output = data.get("console_output", "") or ""

    # Recent real errors, pulled from the actual deployment console log.
    # Excludes routine Docker Compose "orphan containers" chatter and the
    # expected Mongo "duplicate key" noise from re-registering an existing
    # subscriber — neither indicates an actual problem.
    NOISE_SUBSTRINGS = ["orphan containers", "level=warning", "duplicate key error"]
    error_lines = []
    for ln in console_output.splitlines():
        lower = ln.lower()
        if "error" not in lower and "failed" not in lower:
            continue
        if any(noise in lower for noise in NOISE_SUBSTRINGS):
            continue
        error_lines.append(ln.strip())
    seen = set()
    recent_errors = []
    for ln in error_lines:
        if ln and ln not in seen:
            seen.add(ln)
            recent_errors.append(ln)
    recent_errors = recent_errors[-20:]

    # Prometheus scrape target health via the generic 'up' metric — this
    # works regardless of exporter naming, so it's always a safe real check.
    prom_targets = []
    try:
        presp = requests.get(f"{PROMETHEUS_URL}/api/v1/query", params={"query": "up"}, timeout=8)
        if presp.ok and presp.json().get("status") == "success":
            for series in presp.json()["data"]["result"]:
                metric = series.get("metric", {})
                job = metric.get("job", "unknown")
                instance = metric.get("instance", "")
                val = series.get("value", [None, None])[1]
                prom_targets.append((job, instance, "UP" if val == "1" else "DOWN"))
    except Exception:
        pass

    # Real KPI values — discover actual Open5GS metric names from Prometheus
    # itself (never guessed/invented), then query current values for a
    # bounded set of them. This is Python doing the discovery, not the LLM,
    # so there's no risk of a fabricated metric name reaching the report.
    kpi_rows = []
    try:
        names_resp = requests.get(f"{PROMETHEUS_URL}/api/v1/label/__name__/values", timeout=8)
        if names_resp.ok:
            all_names = names_resp.json().get("data", [])
            relevant = [n for n in all_names if n.startswith("fivegs_") or "ue" in n.lower() or "session" in n.lower()]
            for metric_name in sorted(relevant)[:25]:
                try:
                    qresp = requests.get(f"{PROMETHEUS_URL}/api/v1/query", params={"query": metric_name}, timeout=5)
                    if qresp.ok and qresp.json().get("status") == "success":
                        results = qresp.json()["data"]["result"]
                        if results:
                            val = results[0].get("value", [None, None])[1]
                            kpi_rows.append((metric_name, val))
                except Exception:
                    continue
    except Exception:
        pass

    def _sorted_items(d):
        try:
            return sorted(d.items(), key=lambda kv: int(kv[0]))
        except Exception:
            return sorted(d.items())

    # Brief per-device activity summary — capped so this stays fast and
    # readable even at scale (a full log dump per device would bloat the
    # report and slow generation on larger deployments).
    gnb_summaries = []
    for gid, ginfo in _sorted_items(gnbs)[:MAX_DEVICES_FOR_LOG_SUMMARY]:
        if ginfo.get("status") != "running":
            continue
        raw = _fetch_container_logs(f"nr_gnb{gid}")
        gnb_summaries.append((gid, _summarize_log(raw, LOG_KEYWORDS_GNB)))

    ue_summaries = []
    for uid, uinfo in _sorted_items(ues)[:MAX_DEVICES_FOR_LOG_SUMMARY]:
        if uinfo.get("status") != "running":
            continue
        raw = _fetch_container_logs(f"nr_ue{uid}")
        ue_summaries.append((uid, _summarize_log(raw, LOG_KEYWORDS_UE)))

    doc = Document()
    doc.add_heading(clean_text("5G Network Status Report"), level=0)
    doc.add_paragraph(clean_text(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}"))

    doc.add_heading(clean_text("Core Network"), level=1)
    doc.add_paragraph(clean_text(f"State: {core_state}"))
    doc.add_paragraph(clean_text("Functions: AMF, SMF, UPF, UDM, UDR, AUSF, NRF"))

    amf_running = [k for k, v in amfs.items() if v.get("status") == "running"]
    amf_error = [k for k, v in amfs.items() if v.get("status") != "running"]
    doc.add_heading(clean_text(f"AMF Instances ({len(amf_running)} up, {len(amf_error)} error)"), level=2)
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    t.rows[0].cells[0].text = clean_text("Instance")
    t.rows[0].cells[1].text = clean_text("Status")
    t.rows[0].cells[2].text = clean_text("IP")
    for k, v in _sorted_items(amfs):
        row = t.add_row().cells
        row[0].text = clean_text(str(k))
        row[1].text = clean_text(str(v.get("status", "unknown")))
        row[2].text = clean_text(str(v.get("ip") or "-"))

    doc.add_heading(clean_text(f"gNodeBs ({len(gnbs)})"), level=1)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(["gNB ID", "TAC", "IP", "Status"]):
        t.rows[0].cells[i].text = clean_text(h)
    for k, v in _sorted_items(gnbs):
        row = t.add_row().cells
        row[0].text = clean_text(str(k))
        row[1].text = clean_text(str(v.get("tac", "-")))
        row[2].text = clean_text(str(v.get("ip", "-")))
        row[3].text = clean_text(str(v.get("status", "unknown")))

    doc.add_heading(clean_text(f"Connected UEs ({len(ues)})"), level=1)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(["UE ID", "gNB", "IP", "Status"]):
        t.rows[0].cells[i].text = clean_text(h)
    for k, v in _sorted_items(ues):
        row = t.add_row().cells
        row[0].text = clean_text(str(k))
        row[1].text = clean_text(str(v.get("gnb", "-")))
        row[2].text = clean_text(str(v.get("ip", "-")))
        row[3].text = clean_text(str(v.get("status", "unknown")))

    doc.add_heading(clean_text("Prometheus Scrape Targets"), level=1)
    if prom_targets:
        t = doc.add_table(rows=1, cols=3)
        t.style = "Light Grid Accent 1"
        for i, h in enumerate(["Job", "Instance", "Status"]):
            t.rows[0].cells[i].text = clean_text(h)
        for job, instance, status in prom_targets:
            row = t.add_row().cells
            row[0].text = clean_text(job)
            row[1].text = clean_text(instance)
            row[2].text = clean_text(status)
    else:
        doc.add_paragraph(clean_text("Prometheus target data was not available at generation time."))

    doc.add_heading(clean_text(f"Live KPI Metrics ({len(kpi_rows)})"), level=1)
    if kpi_rows:
        t = doc.add_table(rows=1, cols=2)
        t.style = "Light Grid Accent 1"
        t.rows[0].cells[0].text = clean_text("Metric")
        t.rows[0].cells[1].text = clean_text("Current Value")
        for name, val in kpi_rows:
            row = t.add_row().cells
            row[0].text = clean_text(name)
            row[1].text = clean_text(str(val))
    else:
        doc.add_paragraph(
            clean_text(
                "No 5G-specific KPI metrics were found on Prometheus at generation time "
                "(no series matching Open5GS's metric naming were returned)."
            )
        )

    doc.add_heading(clean_text("Per-Device Activity Summary"), level=1)
    doc.add_paragraph(
        clean_text(
            "Milestone lines (registration, session establishment, failures) pulled directly "
            "from each running device's live container log."
        )
    )

    doc.add_heading(clean_text("gNodeBs"), level=2)
    if gnb_summaries:
        for gid, lines in gnb_summaries:
            doc.add_heading(clean_text(f"gNB {gid}"), level=3)
            if lines:
                for ln in lines:
                    doc.add_paragraph(clean_text(ln), style="List Bullet")
            else:
                doc.add_paragraph(clean_text("No matching milestone lines found in the current log."))
    else:
        doc.add_paragraph(clean_text("No running gNodeBs to summarize."))

    doc.add_heading(clean_text("UEs"), level=2)
    if ue_summaries:
        for uid, lines in ue_summaries:
            doc.add_heading(clean_text(f"UE {uid}"), level=3)
            if lines:
                for ln in lines:
                    doc.add_paragraph(clean_text(ln), style="List Bullet")
            else:
                doc.add_paragraph(clean_text("No matching milestone lines found in the current log."))
    else:
        doc.add_paragraph(clean_text("No running UEs to summarize."))

    doc.add_heading(clean_text(f"Recent Issues ({len(recent_errors)})"), level=1)
    if recent_errors:
        for ln in recent_errors:
            doc.add_paragraph(clean_text(ln), style="List Bullet")
    else:
        doc.add_paragraph(clean_text("No errors found in the recent deployment log."))

    filename = f"network_report_{time.strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = REPORTS_DIR / filename
    doc.save(str(filepath))

    download_url = f"http://localhost:5001/download-report/{filename}"
    # Also sanitize the summary string in case any of the numbers have control chars (unlikely)
    summary = (
        f"Report generated: {filename}\n"
        f"Core: {core_state} | AMFs: {len(amf_running)} up, {len(amf_error)} error | "
        f"gNBs: {len(gnbs)} | UEs: {len(ues)} | KPIs found: {len(kpi_rows)} | "
        f"Device summaries: {len(gnb_summaries)} gNBs, {len(ue_summaries)} UEs | "
        f"Issues logged: {len(recent_errors)}\n"
        f"Download: {download_url}"
    )
    return clean_text(summary)

def execute_action(action_name, payload=None):
    """Sends the command payload safely to your main app.py orchestrator."""
    if payload is None:
        payload = {}
    payload["action"] = action_name
    try:
        response = requests.post(ORCHESTRATOR_URL, json=payload, timeout=10)
        if response.ok:
            return f"Success: {response.json()}"
        return f"Orchestrator error: {response.status_code} - {response.text}"
    except requests.RequestException as e:
        return f"Failed to connect to Flask app at {ORCHESTRATOR_URL}: {e}"

# ---------------------------------------------------------
# KPI tools: Prometheus (live values) + Grafana (KPI discovery)
# ---------------------------------------------------------
def _grafana_headers():
    headers = {"Accept": "application/json"}
    if GRAFANA_API_KEY:
        headers["Authorization"] = f"Bearer {GRAFANA_API_KEY}"
    return headers

def query_prometheus(promql: str, **kwargs) -> str:
    """Run a PromQL instant query against Prometheus to get a live KPI value.
    Use this whenever the user asks for a specific number: registered UEs, PDU
    session counts, throughput, latency, error rates, container-level metrics, etc.
    If you're unsure of the exact metric name, call list_prometheus_metrics or
    get_grafana_dashboard_kpis first to find the correct expression.
    Args:
        promql: The PromQL expression, e.g. 'sum(rate(fivegs_amfstatssubinfo_amf_registered_ue{}[5m]))'
    """
    try:
        resp = requests.get(f"{PROMETHEUS_URL}/api/v1/query", params={"query": promql}, timeout=8)
        if not resp.ok:
            return f"Prometheus query failed: {resp.status_code} {resp.text[:300]}"
        data = resp.json()
        if data.get("status") != "success":
            return f"Prometheus error: {data}"
        result = data["data"]["result"]
        if not result:
            return "Query returned no data (metric may not exist yet, or has no samples)."
        lines = []
        for series in result[:20]:
            metric = series.get("metric", {})
            value = series.get("value", [None, None])[1]
            label = ", ".join(f"{k}={v}" for k, v in metric.items()) or "value"
            lines.append(f"{label}: {value}")
        return "\n".join(lines)
    except Exception as e:
        return f"Failed to query Prometheus at {PROMETHEUS_URL}: {e}"

def list_prometheus_metrics(filter_text: str = "", **kwargs) -> str:
    """List available Prometheus metric names, optionally filtered by a substring.
    Use this to discover what KPIs actually exist before guessing a metric name.
    Args:
        filter_text: Optional substring to filter metric names (e.g. 'amf', 'ue', 'gnb', 'ran')
    """
    try:
        resp = requests.get(f"{PROMETHEUS_URL}/api/v1/label/__name__/values", timeout=8)
        if not resp.ok:
            return f"Failed to list metrics: {resp.status_code}"
        names = resp.json().get("data", [])
        if filter_text:
            names = [n for n in names if filter_text.lower() in n.lower()]
        names = sorted(names)[:80]
        return "\n".join(names) if names else "No matching metrics found."
    except Exception as e:
        return f"Failed to reach Prometheus at {PROMETHEUS_URL}: {e}"

def search_grafana_dashboards(query: str = "", **kwargs) -> str:
    """Search Grafana for dashboards by title, to find which dashboard covers a KPI area
    (e.g. 'AMF', 'throughput', 'handover'). Returns titles and uids.
    Args:
        query: Search text to match dashboard titles (empty returns all dashboards)
    """
    try:
        resp = requests.get(f"{GRAFANA_URL}/api/search", params={"query": query}, headers=_grafana_headers(), timeout=8)
        if resp.status_code == 401:
            return "Grafana returned 401 Unauthorized. GRAFANA_API_KEY needs to be set (Service Account token)."
        if not resp.ok:
            return f"Grafana search failed: {resp.status_code} {resp.text[:300]}"
        dashboards = resp.json()
        if not dashboards:
            return "No dashboards found."
        return "\n".join(f"{d.get('title')} (uid={d.get('uid')})" for d in dashboards[:30])
    except Exception as e:
        return f"Failed to reach Grafana at {GRAFANA_URL}: {e}"

def get_grafana_dashboard_kpis(dashboard_uid: str, **kwargs) -> str:
    """Fetch a Grafana dashboard's panels and the exact PromQL expression behind each one.
    Use this after search_grafana_dashboards to find the right expression for a named KPI,
    then pass that expression to query_prometheus to get the live value.
    Args:
        dashboard_uid: The dashboard uid returned by search_grafana_dashboards
    """
    try:
        resp = requests.get(f"{GRAFANA_URL}/api/dashboards/uid/{dashboard_uid}", headers=_grafana_headers(), timeout=8)
        if resp.status_code == 401:
            return "Grafana returned 401 Unauthorized. GRAFANA_API_KEY needs to be set (Service Account token)."
        if not resp.ok:
            return f"Failed to fetch dashboard: {resp.status_code} {resp.text[:300]}"
        model = resp.json().get("dashboard", {})
        lines = []
        for panel in model.get("panels", []):
            title = panel.get("title", "Untitled panel")
            targets = panel.get("targets", [])
            exprs = [t.get("expr") for t in targets if t.get("expr")]
            if exprs:
                lines.append(f"{title}: " + " | ".join(exprs))
        return "\n".join(lines) if lines else "No PromQL-backed panels found on this dashboard."
    except Exception as e:
        return f"Failed to reach Grafana at {GRAFANA_URL}: {e}"

# ---------------------------------------------------------
# Registries (defined after all functions to avoid NameError)
# ---------------------------------------------------------
TOOL_REGISTRY = {
    "start_core": start_core,
    "stop_core": stop_core,
    "start_gnb": start_gnb,
    "add_ue": add_ue,
    "get_ue_logs": get_ue_logs,
    "get_gnb_logs": get_gnb_logs,
    "get_core_logs": get_core_logs,
    "stop_all": stop_all,
    "get_logs_or_files": get_logs_or_files,
    "read_yaml_file": read_yaml_file,
    "check_container_health": check_container_health,
    "launch_full_network": launch_full_network,
    "get_network_topology": get_network_topology,
    "generate_network_report": generate_network_report,
    "query_prometheus": query_prometheus,
    "list_prometheus_metrics": list_prometheus_metrics,
    "search_grafana_dashboards": search_grafana_dashboards,
    "get_grafana_dashboard_kpis": get_grafana_dashboard_kpis,
}

available_tools = list(TOOL_REGISTRY.values())

def new_conversation():
    return [{'role': 'system', 'content': SYSTEM_PROMPT}]

# ---------------------------------------------------------
# Agentic loop: resolve tool calls (possibly chained), used by
# both CLI and web modes. Mutates and returns `messages`.
# ---------------------------------------------------------
def resolve_tool_calls(messages, max_iterations=MAX_TOOL_ITERATIONS, on_tool_call=None):
    """Runs non-streaming ollama.chat calls, executing any requested tools,
    looping up to max_iterations so the agent can chain tools (e.g. search a
    Grafana dashboard, then query Prometheus with what it found).
    on_tool_call(name, args, result) is called after each tool execution, if provided.

    Returns (messages, final_text). final_text is the model's actual answer,
    taken from the SAME call that decided no more tools were needed — we do
    NOT make a second call to "re-ask" for the answer, because by that point
    the conversation already ends in an assistant turn and asking the model
    to speak again right after itself is confusing and often yields nothing.
    """
    final_text = ""
    for _ in range(max_iterations):
        response = ollama.chat(model=OLLAMA_MODEL, messages=messages, tools=available_tools)
        msg = response['message']
        messages.append(msg)

        tool_calls = msg.get('tool_calls')
        if not tool_calls:
            final_text = msg.get('content', '') or ""
            break

        for tool in tool_calls:
            func_name = tool['function']['name']
            func_args = tool['function']['arguments'] or {}

            if func_name in TOOL_REGISTRY:
                try:
                    result = TOOL_REGISTRY[func_name](**func_args)
                except Exception as e:
                    result = f"Tool '{func_name}' raised an error: {e}"
            else:
                result = f"Unknown tool requested: {func_name}"

            if on_tool_call:
                on_tool_call(func_name, func_args, result)

            messages.append({'role': 'tool', 'content': str(result), 'name': func_name})

    if not final_text:
        final_text = (
            "I ran the requested step(s) above. Let me know if you'd like the details "
            "or want me to check on anything else."
        )
        messages.append({'role': 'assistant', 'content': final_text})

    return messages, final_text

# ---------------------------------------------------------
# Flask Web App (standalone, port 5001)
# ---------------------------------------------------------
web_app = Flask(__name__)
web_conversations = {}

@web_app.route("/")
def web_index():
    return render_template("ai_chat.html")

@web_app.route("/api/ai/chat", methods=["POST"])
def api_ai_chat():
    data = request.json or {}
    user_message = (data.get("message") or "").strip()
    session_id = data.get("session_id") or "default"

    if not user_message:
        return jsonify({"status": "error", "message": "Empty message."}), 400

    if session_id not in web_conversations:
        web_conversations[session_id] = new_conversation()

    messages = web_conversations[session_id]
    messages.append({'role': 'user', 'content': user_message})

    def generate_events():
        def on_tool_call(name, args, result):
            print(f"[TOOL RESULT] {name}({args}) -> {str(result)[:300]}", flush=True)

        tool_events = []

        def collect_tool_call(name, args, result):
            on_tool_call(name, args, result)
            tool_events.append({'type': 'tool', 'name': name})
            if name == 'get_network_topology':
                try:
                    tool_events.append({'type': 'topology', 'data': json.loads(result)})
                except Exception:
                    pass
            if name == 'generate_network_report':
                url_match = re.search(r'(https?://\S+\.docx)', str(result))
                if url_match:
                    url = url_match.group(1)
                    filename = url.rsplit('/', 1)[-1]
                    tool_events.append({'type': 'report', 'url': url, 'filename': filename})

        _, final_text = resolve_tool_calls(messages, on_tool_call=collect_tool_call)
        print(f"[FINAL REPLY] {final_text[:300]}", flush=True)

        for evt in tool_events:
            yield f"data: {json.dumps(evt)}\n\n"

        # Simulate a typing effect over the already-generated text (word by
        # word) rather than making a second, confusing call to the model.
        for piece in re.findall(r'\S+\s*', final_text):
            yield f"data: {json.dumps({'type': 'token', 'text': piece})}\n\n"
            time.sleep(0.02)

        yield f"data: {json.dumps({'type': 'done'})}\n\n"

    return Response(generate_events(), mimetype='text/event-stream')

@web_app.route("/download-report/<path:filename>")
def download_report(filename):
    return send_from_directory(str(REPORTS_DIR.resolve()), filename, as_attachment=True)

@web_app.route("/api/ai/reset", methods=["POST"])
def api_ai_reset():
    data = request.json or {}
    session_id = data.get("session_id") or "default"
    web_conversations.pop(session_id, None)
    return jsonify({"status": "reset"})

# ---------------------------------------------------------
# CLI Mode
# ---------------------------------------------------------
def run_ai_cli():
    print("==================================================")
    print("🌐 NetSage AI - 5G Standalone Orchestrator (Ollama)")
    print("Make sure your main app.py is running on port 5000!")
    print("Type commands like: 'Check ue 1 logs' or 'Get amf logs'")
    print("Type 'exit' to quit.")
    print("==================================================")

    messages = new_conversation()

    while True:
        try:
            user_input = input("\nUser Command > ")
            if user_input.lower() in ['exit', 'quit']:
                print("Exiting NetSage AI.")
                break
            if not user_input.strip():
                continue

            messages.append({'role': 'user', 'content': user_input})

            def _log_tool(name, args, result):
                print(f"\n[🔧 tool] {name}({args}) -> {str(result)[:200]}")

            with Spinner("NetSage AI is analyzing your request..."):
                _, final_text = resolve_tool_calls(messages, on_tool_call=_log_tool)

            print(f"\n[NetSage AI]: {final_text}")

        except Exception as e:
            print(f"\n[Error]: {e}")

if __name__ == "__main__":
    if len(sys.argv) > 1 and sys.argv[1] == "--web":
        print("[🌐 NetSage AI Web Server] Running on http://0.0.0.0:5001")
        web_app.run(host="0.0.0.0", port=5001, debug=True, threaded=True)
    else:
        run_ai_cli()
